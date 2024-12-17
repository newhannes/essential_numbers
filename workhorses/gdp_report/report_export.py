from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

VISUALS_PATH = "output/charts"


#### ==== SECTION 5: EXPORTING TO WORD ==== ####
## MARK: Exporting to Word
def generate_word_doc(OVERVIEW_TEXT, CONTRIBUTOR_TEXT, CHANGES_TEXT, REVISION_TEXT):
    # Create a new Document
    doc = Document()

    # Add Title
    doc.add_heading('GDP Report', level=1)

    # Add Overview Section
    doc.add_heading('Overview', level=2)
    doc.add_paragraph(OVERVIEW_TEXT)
    doc.add_picture(f'{VISUALS_PATH}/since22.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Contributors Section
    doc.add_heading('Contributors to GDP Growth', level=2)
    doc.add_paragraph(CONTRIBUTOR_TEXT)
    doc.add_picture(f'{VISUALS_PATH}/contributors.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Changes This Quarter Section
    doc.add_heading('Changes This Quarter', level=2)
    doc.add_paragraph(CHANGES_TEXT)
    doc.add_picture(f'{VISUALS_PATH}/growth_comparison.png', width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Revisions Section
    if REVISION_TEXT:
        doc.add_heading('Revisions', level=2)
        doc.add_paragraph(REVISION_TEXT)
        doc.add_picture(f'{VISUALS_PATH}/revisions.png', width=Inches(6))

    # Save the Document
    doc.save('output/GDP REPORT 12.13.docx')

#### ==== SECTION 6: EXPORTING TO HTML ==== ####
## MARK: Exporting to HTML
def generate_html_report(OVERVIEW_TEXT, CONTRIBUTOR_TEXT, CHANGES_TEXT, REVISION_TEXT):

    if REVISION_TEXT:
        revision_html = f"""<section>
                    <h2>Revisions</h2>
                    <p>{REVISION_TEXT}</p>
                    <div class="table">
                        <img src="/mount/src/essential_numbers/output/charts/revisions.png" alt="Revisions to Growth in 3rd Estimate">
                    </div>
                </section>"""
    else:
        revision_html = ""

    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>GDP Report</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                margin: 0;
                padding: 0;
                background-color: #f4f4f4;
            }}
            .container {{
                width: 80%;
                margin: auto;
                overflow: hidden;
            }}
            header {{
                background: #333;
                color: #fff;
                padding-top: 30px;
                min-height: 70px;
                text-align: center;
                width: 70%;
                margin: 0 auto 
            }}
            header h1 {{
                text-align: center;
                text-transform: uppercase;
                margin: 0;
                font-size: 24px;
            }}
            section {{
                padding: 20px;
                margin: 20px 0;
                background: #fff;
                border-radius: 10px;
            }}
            section h2 {{
                text-align: left;
                color: #333;
            }}
            section p {{
                line-height: 1.6;
            }}
            .chart {{
                text-align: left;
                margin: 20px 0;
            }}
            .chart img {{
                max-width: 80%;
                height: auto;
            }}
            .table {{
                text-align: center;
                margin: 20px 0;
            }}
            .table img {{
                max-width: 80%;
                height: auto;
            }}
        </style>
    </head>
    <body>
        <header>
            <h1>GDP Report</h1>
        </header>
        <div class="container">
            <section>
                <h2>Overview</h2>
                <p>{OVERVIEW_TEXT}</p>
                <div class="chart">
                    <img src="/mount/src/essential_numbers/output/charts/since22.png" alt="Annualized real GDP growth since Q1 2022">
                </div>
            </section>
            <section>
                <h2>Contributors to GDP Growth</h2>
                <p>{CONTRIBUTOR_TEXT}</p>
                <div class="chart">
                    <img src="/mount/src/essential_numbers/output/charts/contributors.png" alt="Percentage point contributions to GDP growth">
                </div>
            </section>
            <section>
                <h2>Changes This Quarter</h2>
                <p>{CHANGES_TEXT}</p>
                <div class="chart">
                    <img src="/mount/src/essential_numbers/output/charts/growth_comparison.png" alt="Growth in GDP components in Q3 2024">
                </div>
            </section>
            {revision_html}
        </div>
    </body>
    </html>
    """
    
    # Save the HTML content to a file
    with open("output/GDP_REPORT.html", "w") as file:
        file.write(html_content)

    return html_content